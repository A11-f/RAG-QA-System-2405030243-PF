from docx import Document
import os
import config


def create_sample_doc1():
    doc = Document()
    doc.add_heading('自然语言处理简介', 0)
    
    doc.add_heading('什么是自然语言处理？', level=1)
    doc.add_paragraph(
        '自然语言处理（Natural Language Processing，简称NLP）是人工智能（AI）的一个分支，'
        '专注于计算机与人类（自然）语言之间的交互。它是计算机科学、人工智能和语言学的交叉领域。'
    )
    
    doc.add_heading('NLP的主要任务', level=1)
    doc.add_paragraph('NLP包含多种任务，包括但不限于：')
    doc.add_paragraph('1. 文本分类：将文本分为预定义的类别', style='List Number')
    doc.add_paragraph('2. 命名实体识别：识别人名、地名、组织机构名等', style='List Number')
    doc.add_paragraph('3. 机器翻译：将一种语言翻译成另一种语言', style='List Number')
    doc.add_paragraph('4. 情感分析：分析文本的情感倾向（正面、负面、中性）', style='List Number')
    doc.add_paragraph('5. 问答系统：自动回答用户问题', style='List Number')
    
    doc.add_heading('应用场景', level=1)
    doc.add_paragraph('NLP技术广泛应用于：搜索引擎、智能助手、自动客服、内容推荐系统、文本摘要等领域。')
    
    doc.save(os.path.join(config.DATA_DIR, '自然语言处理简介.docx'))


def create_sample_doc2():
    doc = Document()
    doc.add_heading('词向量与词嵌入', 0)
    
    doc.add_heading('词向量的概念', level=1)
    doc.add_paragraph(
        '词向量（Word Embeddings）是将词语映射到实数向量的技术，'
        '是NLP中最基础且最重要的技术之一。'
    )
    
    doc.add_heading('Word2Vec', level=1)
    doc.add_paragraph(
        'Word2Vec是Google在2013年提出的词向量训练方法，包含两种模型：'
    )
    doc.add_paragraph('1. CBOW（Continuous Bag-of-Words）：根据上下文预测中心词', style='List Number')
    doc.add_paragraph('2. Skip-gram：根据中心词预测上下文', style='List Number')
    
    doc.add_heading('词向量的应用', level=1)
    doc.add_paragraph('词向量广泛应用于文本分类、情感分析、命名实体识别等任务中。')
    
    doc.save(os.path.join(config.DATA_DIR, '词向量与词嵌入.docx'))


def create_sample_doc3():
    doc = Document()
    doc.add_heading('Transformer模型结构详解', 0)
    
    doc.add_heading('Transformer概述', level=1)
    doc.add_paragraph(
        'Transformer是Google在2017年发表的论文《Attention Is All You Need》中提出的模型架构，'
        '完全基于注意力机制，不使用RNN或CNN。'
    )
    
    doc.add_heading('核心组件', level=1)
    doc.add_paragraph('Transformer主要由以下组件构成：')
    doc.add_paragraph('1. 自注意力机制（Self-Attention）', style='List Number')
    doc.add_paragraph('2. 多头注意力（Multi-Head Attention）', style='List Number')
    doc.add_paragraph('3. 位置编码（Positional Encoding）', style='List Number')
    doc.add_paragraph('4. 前馈网络（Feed-Forward Network）', style='List Number')
    doc.add_paragraph('5. 层归一化和残差连接', style='List Number')
    
    doc.add_heading('架构', level=1)
    doc.add_paragraph('Transformer包含编码器（Encoder）和解码器（Decoder）两部分。')
    
    doc.save(os.path.join(config.DATA_DIR, 'Transformer模型结构.docx'))


def create_sample_doc4():
    doc = Document()
    doc.add_heading('BERT模型介绍', 0)
    
    doc.add_heading('什么是BERT', level=1)
    doc.add_paragraph(
        'BERT（Bidirectional Encoder Representations from Transformers）'
        '是Google在2018年提出的预训练语言模型。'
    )
    
    doc.add_heading('BERT的特点', level=1)
    doc.add_paragraph('1. 双向上下文：同时考虑词语的左右上下文', style='List Number')
    doc.add_paragraph('2. 预训练+微调：两阶段训练流程', style='List Number')
    doc.add_paragraph('3. 预训练任务：Masked Language Model和Next Sentence Prediction', style='List Number')
    
    doc.add_heading('应用', level=1)
    doc.add_paragraph('BERT在文本分类、问答系统、命名实体识别等任务上都取得了优异的表现。')
    
    doc.save(os.path.join(config.DATA_DIR, 'BERT模型介绍.docx'))


def create_sample_doc5():
    doc = Document()
    doc.add_heading('文本分类技术', 0)
    
    doc.add_heading('文本分类简介', level=1)
    doc.add_paragraph('文本分类是将文本分配到预定义类别的任务，是NLP的基础任务之一。')
    
    doc.add_heading('传统方法', level=1)
    doc.add_paragraph('1. 词袋模型（Bag-of-Words）', style='List Number')
    doc.add_paragraph('2. TF-IDF', style='List Number')
    doc.add_paragraph('3. SVM、朴素贝叶斯等', style='List Number')
    
    doc.add_heading('深度学习方法', level=1)
    doc.add_paragraph('1. TextCNN', style='List Number')
    doc.add_paragraph('2. LSTM/GRU', style='List Number')
    doc.add_paragraph('3. BERT等预训练模型', style='List Number')
    
    doc.add_heading('流程', level=1)
    doc.add_paragraph('一般流程：数据预处理 → 特征提取 → 模型训练 → 评估。')
    
    doc.save(os.path.join(config.DATA_DIR, '文本分类技术.docx'))


if __name__ == '__main__':
    print('正在创建示例文档...')
    create_sample_doc1()
    create_sample_doc2()
    create_sample_doc3()
    create_sample_doc4()
    create_sample_doc5()
    print('✅ 示例文档创建完成！')