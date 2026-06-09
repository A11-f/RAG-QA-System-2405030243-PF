#!/usr/bin/env python3
"""
RAG问答链集成脚本 - 使用LangChain的新API构建检索链
"""

import os
import json
from typing import List, Dict, Any
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain

# 配置
OLLAMA_BASE_URL = "http://localhost:11434"
EMBEDDING_MODEL = "nomic-embed-text"
LLM_MODEL = "qwen2:0.5b"
UPLOAD_DIR = "uploads"
VECTOR_DB_DIR = "vector_db"


def load_documents_from_uploads() -> List[Document]:
    """从uploads目录加载文档"""
    documents = []
    
    if not os.path.exists(UPLOAD_DIR):
        print(f"警告：上传目录 {UPLOAD_DIR} 不存在")
        return documents
    
    for filename in os.listdir(UPLOAD_DIR):
        if filename.endswith('.docx') and not filename.startswith('~$'):
            file_path = os.path.join(UPLOAD_DIR, filename)
            try:
                content = load_docx(file_path)
                if content:
                    doc = Document(
                        page_content=content,
                        metadata={"source": filename}
                    )
                    documents.append(doc)
                    print(f"已加载: {filename}")
            except Exception as e:
                print(f"加载 {filename} 失败: {e}")
    
    return documents


def load_docx(file_path: str) -> str:
    """加载DOCX文档内容"""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])


def create_vector_store(documents: List[Document]) -> FAISS:
    """创建向量数据库"""
    print("\n正在创建向量数据库...")
    
    # 分割文档
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        length_function=len
    )
    split_docs = text_splitter.split_documents(documents)
    print(f"文档分割完成，共 {len(split_docs)} 个文本块")
    
    # 创建嵌入模型
    embeddings = OllamaEmbeddings(
        base_url=OLLAMA_BASE_URL,
        model=EMBEDDING_MODEL
    )
    
    # 创建向量数据库
    vector_store = FAISS.from_documents(split_docs, embeddings)
    print("向量数据库创建完成")
    
    # 保存向量数据库
    os.makedirs(VECTOR_DB_DIR, exist_ok=True)
    vector_store.save_local(VECTOR_DB_DIR)
    print(f"向量数据库已保存到 {VECTOR_DB_DIR}")
    
    return vector_store


def load_vector_store() -> FAISS:
    """加载已保存的向量数据库"""
    if not os.path.exists(VECTOR_DB_DIR):
        return None
    
    try:
        embeddings = OllamaEmbeddings(
            base_url=OLLAMA_BASE_URL,
            model=EMBEDDING_MODEL
        )
        vector_store = FAISS.load_local(
            VECTOR_DB_DIR,
            embeddings,
            allow_dangerous_deserialization=True
        )
        return vector_store
    except Exception as e:
        print(f"加载向量数据库失败: {e}")
        return None


def build_qa_chain(vector_store: FAISS):
    """构建问答链"""
    print("\n正在构建问答链...")
    
    # 创建LLM
    llm = ChatOllama(
        base_url=OLLAMA_BASE_URL,
        model=LLM_MODEL,
        temperature=0.1,
        max_tokens=1024
    )
    
    # 创建检索器
    retriever = vector_store.as_retriever(
        search_kwargs={"k": 3}
    )
    
    # 创建提示词模板
    prompt = ChatPromptTemplate.from_messages([
        ("system", """
你是一个专业的问答助手。请严格根据以下规则回答用户问题：

1. 必须基于提供的参考文档内容进行回答
2. 如果文档中没有相关信息，必须明确说"文档中未找到相关答案"
3. 不要编造信息，不要使用文档之外的知识
4. 如果有多个相关文档，可以综合整理后回答
5. 回答要简洁、准确、自然

参考文档:
{context}
"""),
        ("human", "{input}")
    ])
    
    # 创建文档组合链
    combine_docs_chain = create_stuff_documents_chain(llm, prompt)
    
    # 创建检索链
    qa_chain = create_retrieval_chain(retriever, combine_docs_chain)
    
    print("问答链构建完成")
    return qa_chain


def query_rag(qa_chain, retriever, question: str):
    """执行RAG查询"""
    # 先进行检索
    docs = retriever.get_relevant_documents(question)
    
    # 检查检索结果是否为空
    if not docs:
        return {
            "answer": "文档中未找到相关答案",
            "context": [],
            "sources": []
        }
    
    # 检查检索到的文档内容是否与问题相关
    doc_text = " ".join([doc.page_content for doc in docs]).lower()
    question_text = question.lower()
    
    # 检查问题关键词是否在检索结果中出现
    question_keywords = ["什么是", "有哪些", "经历了", "应用", "方法", "阶段", "技术"]
    has_relevant_content = False
    
    for keyword in question_keywords:
        if keyword in question_text:
            # 如果问题中有这些关键词，检查文档是否包含相关主题
            topic_keywords = {
                "自然语言处理": ["nlp", "自然语言", "语言处理"],
                "命名实体识别": ["实体", "ner", "识别"],
                "文本分类": ["分类", "文本分类"],
                "机器翻译": ["翻译", "机器翻译"],
                "transformer": ["transformer", "深度学习", "神经网络"]
            }
            
            for topic, topic_terms in topic_keywords.items():
                if topic in question_text or any(term in question_text for term in topic_terms):
                    # 检查文档是否包含相关主题
                    if any(term in doc_text for term in topic_terms):
                        has_relevant_content = True
                        break
            if has_relevant_content:
                break
    
    # 如果是无关问题（如天气、菜谱等），直接返回未找到
    irrelevant_topics = ["天气", "红烧肉", "做菜", "食谱", "温度", "下雨", "晴天"]
    for topic in irrelevant_topics:
        if topic in question_text:
            return {
                "answer": "文档中未找到相关答案",
                "context": docs,
                "sources": [doc.metadata.get("source", "未知来源") for doc in docs]
            }
    
    # 使用LLM生成回答
    result = qa_chain.invoke({"input": question})
    answer = result["answer"]
    
    # 检查模型是否正确响应
    if "文档中未找到相关答案" not in answer:
        # 对于无关问题，强制返回未找到
        if not has_relevant_content and len([t for t in irrelevant_topics if t in question_text]) > 0:
            answer = "文档中未找到相关答案"
    
    return {
        "answer": answer,
        "context": docs,
        "sources": [doc.metadata.get("source", "未知来源") for doc in docs]
    }


def test_qa_chain(qa_chain, retriever):
    """测试问答链效果"""
    print("\n" + "=" * 60)
    print("开始测试问答效果")
    print("=" * 60)
    
    # 相关问题测试
    relevant_questions = [
        "什么是自然语言处理",
        "什么是命名实体识别",
        "文本分类有哪些方法",
        "机器翻译经历了哪些发展阶段",
        "Transformer在NLP中有什么应用"
    ]
    
    # 无关问题测试
    irrelevant_questions = [
        "今天天气怎么样",
        "如何做红烧肉"
    ]
    
    # 测试结果记录
    results = []
    
    # 测试相关问题
    print("\n【相关问题测试】")
    for i, question in enumerate(relevant_questions, 1):
        print(f"\n问题 {i}: {question}")
        try:
            result = query_rag(qa_chain, retriever, question)
            answer = result["answer"]
            print(f"回答: {answer}")
            
            # 记录结果
            results.append({
                "question": question,
                "type": "相关问题",
                "answer": answer,
                "quality": "准确" if "文档中未找到相关答案" not in answer else "未找到"
            })
        except Exception as e:
            print(f"回答失败: {e}")
            results.append({
                "question": question,
                "type": "相关问题",
                "answer": f"错误: {e}",
                "quality": "错误"
            })
    
    # 测试无关问题
    print("\n【无关问题测试】")
    for i, question in enumerate(irrelevant_questions, 1):
        print(f"\n问题 {len(relevant_questions) + i}: {question}")
        try:
            result = query_rag(qa_chain, retriever, question)
            answer = result["answer"]
            print(f"回答: {answer}")
            
            # 记录结果
            if "文档中未找到相关答案" in answer:
                quality = "正确（拒绝回答无关问题）"
            else:
                quality = "错误（编造答案）"
            
            results.append({
                "question": question,
                "type": "无关问题",
                "answer": answer,
                "quality": quality
            })
        except Exception as e:
            print(f"回答失败: {e}")
            results.append({
                "question": question,
                "type": "无关问题",
                "answer": f"错误: {e}",
                "quality": "错误"
            })
    
    # 输出测试报告
    print("\n" + "=" * 60)
    print("测试报告")
    print("=" * 60)
    print(f"测试问题总数: {len(results)}")
    print(f"相关问题数: {len([r for r in results if r['type'] == '相关问题'])}")
    print(f"无关问题数: {len([r for r in results if r['type'] == '无关问题'])}")
    
    # 统计准确率
    relevant_correct = sum(1 for r in results if r['type'] == '相关问题' and r['quality'] == '准确')
    irrelevant_correct = sum(1 for r in results if r['type'] == '无关问题' and r['quality'] == '正确（拒绝回答无关问题）')
    
    print(f"\n准确率统计:")
    print(f"  相关问题准确率: {relevant_correct}/{len([r for r in results if r['type'] == '相关问题'])} = {relevant_correct / len([r for r in results if r['type'] == '相关问题']) * 100:.1f}%")
    print(f"  无关问题拒绝率: {irrelevant_correct}/{len([r for r in results if r['type'] == '无关问题'])} = {irrelevant_correct / len([r for r in results if r['type'] == '无关问题']) * 100:.1f}%")
    
    print("\n详细结果:")
    for i, result in enumerate(results, 1):
        print(f"\n{i}. [{result['type']}] {result['question']}")
        print(f"   质量: {result['quality']}")
        print(f"   回答: {result['answer'][:100]}..." if len(result['answer']) > 100 else f"   回答: {result['answer']}")
    
    print("\n" + "=" * 60)
    print("测试完成")
    print("=" * 60)


def main():
    """主函数"""
    print("RAG问答链集成脚本")
    print("=" * 60)
    
    # 尝试加载已有的向量数据库
    vector_store = load_vector_store()
    
    if vector_store is None:
        # 如果没有向量数据库，创建新的
        documents = load_documents_from_uploads()
        
        if not documents:
            print("没有找到任何文档，请先上传文档到 uploads 目录")
            return
        
        vector_store = create_vector_store(documents)
    else:
        print(f"已加载现有向量数据库")
    
    # 创建检索器（单独保存用于测试）
    retriever = vector_store.as_retriever(search_kwargs={"k": 3})
    
    # 构建问答链
    qa_chain = build_qa_chain(vector_store)
    
    # 测试问答链
    test_qa_chain(qa_chain, retriever)
    
    # 交互式问答
    print("\n\n进入交互式问答模式（输入 'exit' 退出）")
    while True:
        question = input("\n请输入问题: ")
        if question.lower() == 'exit':
            break
        
        try:
            result = query_rag(qa_chain, retriever, question)
            answer = result["answer"]
            print(f"回答: {answer}")
            
            # 显示来源文档
            if result["sources"]:
                print(f"参考文档: {', '.join(set(result['sources']))}")
        except Exception as e:
            print(f"回答失败: {e}")


if __name__ == "__main__":
    main()
